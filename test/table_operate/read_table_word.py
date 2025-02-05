from docx import Document

def list_tables(doc_path):
    """
    读取并打印指定文档的所有表格内容
    """
    doc = Document(doc_path)
    for table_index, table in enumerate(doc.tables):
        print(f"📌 Table {table_index + 1} (共 {len(doc.tables)} 个表格)")
        # 遍历行
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                print(f"  📍 第 {row_index + 1} 行, 第 {col_index + 1} 列: {cell.text.strip()}")
        print("-" * 50)  # 分隔不同表格


def replace_text_preserve_format(cell, replacements):
    """
    替换 Word 表格中的占位符，同时保持原格式
    """
    for para in cell.paragraphs:
        for run in para.runs:
            for key, value in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, value)  # 仅替换内容，不影响格式


def replace_text_in_table(doc, table_index, replacements):
    """
    在指定表格中批量替换占位符
    """
    if table_index >= len(doc.tables):
        raise ValueError(f"表格索引 {table_index} 超出范围，当前文档仅有 {len(doc.tables)} 个表格")

    table = doc.tables[table_index]
    for row in table.rows:
        for cell in row.cells:
            replace_text_preserve_format(cell, replacements)


def delete_table_rows(doc, table_index, start_row, end_row):
    """
    删除 Word 指定表格中的多行
    :param doc: Document 对象
    :param table_index: 需要操作的表格索引 (0-based)
    :param start_row: 删除的起始行索引 (0-based)
    :param end_row: 删除的结束行索引 (0-based)
    """
    if table_index >= len(doc.tables):
        raise ValueError(f"表格索引 {table_index} 超出范围，当前文档仅有 {len(doc.tables)} 个表格")

    table = doc.tables[table_index]
    total_rows = len(table.rows)

    if start_row >= total_rows or end_row >= total_rows:
        raise ValueError(f"表格 {table_index + 1} 仅有 {total_rows} 行，删除范围 ({start_row} - {end_row}) 超出范围")

    # 倒序删除，防止索引变化
    for row_index in range(end_row, start_row - 1, -1):
        table._element.remove(table.rows[row_index]._element)


from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn  # 用于设置字体的 eastAsia 属性


def add_row_with_auto_serial(doc, table_index, cell_values):
    """
    在指定表格末尾添加一行：
      - 第一列自动生成序号（形如“1.”、“2.”）
      - 从第二列开始填充传入的 cell_values
      - 新增的所有文本统一设置为“宋体 + 指定字号”（下面示例用小四 = 12pt）。
    """
    # 1. 检查表格索引
    if table_index >= len(doc.tables):
        raise ValueError(f"表格索引 {table_index} 超出范围，当前文档仅有 {len(doc.tables)} 个表格")

    table = doc.tables[table_index]

    # 2. 新增一行
    new_row = table.add_row()

    # 3. 第一列自动编号（假设表第 1 行是表头，则数据区从第 2 行开始）
    #    如果表没有表头，想让第一条数据就是 "1."，可以直接用 len(table.rows)
    #    若含 1 行表头，则序号 = len(table.rows) - 1（让新增的行从 1. 开始）
    serial_number = len(table.rows) - 1
    new_row.cells[0].text = f"{serial_number}."

    # 4. 给第一列（序号列）设置字体格式
    _apply_font_style(new_row.cells[0], font_name="宋体", font_size_pt=12)

    # 5. 填充后续列内容
    for col_idx in range(1, len(new_row.cells)):
        if (col_idx - 1) < len(cell_values):
            new_row.cells[col_idx].text = str(cell_values[col_idx - 1])
        else:
            new_row.cells[col_idx].text = ""

        # 同样设置为“宋体 + 指定字号”
        _apply_font_style(new_row.cells[col_idx], font_name="宋体", font_size_pt=12)


def _apply_font_style(cell, font_name="宋体", font_size_pt=12):
    """
    工具函数：给单元格内所有段落、所有 run 设置指定字体与字号
    """
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            # 设置西文字体
            run.font.name = font_name
            # 设置中文字体
            run.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            # 设置字号
            run.font.size = Pt(font_size_pt)


def main():
    doc_path = "technical_document_template.docx"
    print("=== 打印所有表格的内容 ===")
    list_tables(doc_path)
    # 1. 打开 Word 文档
    doc_path = "technical_document_template.docx"
    doc = Document(doc_path)

    # 2. 待添加的数据（不含第一列，因为第一列会自动编序号）
    rows_to_add = [
        ["电源滤波组件", "MTLB32B-HNBJ-..."],
        ["保险管", "12.5A"],
        ["导电材料", "/"],
        ["出厂检测报告", "/"],
        ["试验报告", "/"],
        ["纸质合格证", "/"],
    ]

    # 3. 逐行添加数据到第 0 个表格（如果你的目标是第 2 个表格，就改成 table_index=1）
    for row_data in rows_to_add:
        add_row_with_auto_serial(doc, table_index=3, cell_values=row_data)

    # 4. 保存结果
    output_path = "auto_serial_result.docx"
    doc.save(output_path)
    print(f"已在表格中添加新行，并自动编号和设置字体，结果保存为：{output_path}")


if __name__ == "__main__":
    main()

