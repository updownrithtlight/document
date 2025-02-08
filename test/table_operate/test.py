from docx import Document

from app.utils.spec_word_table_processor import SpecWordTableProcessor


def process_specific_table(file_path, selected_items, target_table_index=None, table_header=None):
    """
    处理 Word 文档中的指定表格，根据规则删除未选择项并更新表格内容。

    :param file_path: str, 输入的 Word 文档路径
    :param selected_items: list, 已选择项目的列表
    :param target_table_index: int, 指定表格的索引（从 0 开始），默认为 None
    :param table_header: str, 用于识别目标表格的表头文本，如果提供会优先通过表头匹配表格
    :return: Document 对象，已修改的文档
    """
    # 加载 Word 文档
    doc = Document(file_path)

    # 提取已选择的项目名称
    selected_names = [item['name'] for item in selected_items]

    # Map for replacing True/False to symbols
    boolean_map = {True: '√', False: ''}

    # 目标表格
    target_table = None

    # 定位目标表格
    if table_header:
        # 如果指定了表头，通过表头识别表格
        for table in doc.tables:
            if table.rows[0].cells[0].text.strip() == table_header:
                target_table = table
                break
    elif target_table_index is not None:
        # 如果指定了表格索引，通过索引定位表格
        target_table = doc.tables[target_table_index]

    # 如果没有找到目标表格，抛出异常
    if not target_table:
        raise ValueError("未找到符合条件的表格，请检查参数 `table_header` 或 `target_table_index` 是否正确。")

    # 遍历目标表格并处理数据
    rows_to_delete = []  # 保存需要删除的行
    for row_idx, row in enumerate(target_table.rows[1:], start=1):  # 跳过表头
        cell_name = row.cells[0].text.strip()
        if cell_name not in selected_names:
            # 如果不在已选择项目中，标记为删除
            rows_to_delete.append(row)
        else:
            # 如果在已选择项目中，更新内容
            for item in selected_items:
                if item['name'] == cell_name:
                    row.cells[1].text = boolean_map[item['pcb']]
                    row.cells[2].text = boolean_map[item['beforeSeal']]
                    row.cells[3].text = boolean_map[item['afterLabel']]
                    row.cells[4].text = item['samplePlan']

    # 删除不需要的行
    for row in rows_to_delete:
        target_table._element.remove(row._element)

    return doc


# 示例用法
input_file_path = 'product_specification.docx'
output_file_path = '检验项目_更新.docx'

# 已选择的项目
selected_items = [
    {'project_id': 6, 'key': '1', 'name': '壳体外观', 'pcb': True, 'beforeSeal': False, 'afterLabel': False,
     'samplePlan': '100%'},
    {'project_id': 6, 'key': '2', 'name': '焊点外观', 'pcb': False, 'beforeSeal': True, 'afterLabel': False,
     'samplePlan': '100%'},
    {'project_id': 6, 'key': '6', 'name': '引出端尺寸', 'pcb': False, 'beforeSeal': False, 'afterLabel': False,
     'samplePlan': '100%'},
    {'project_id': 6, 'key': '13', 'name': '插入损耗', 'pcb': False, 'beforeSeal': True, 'afterLabel': False,
     'samplePlan': '100%'},
    {'project_id': 6, 'key': '14', 'name': '常温加电测压降', 'pcb': False, 'beforeSeal': False, 'afterLabel': True,
     'samplePlan': '100%'}
]

# 指定目标表格：通过表头定位或通过索引
table_header = "检验项目"  # 替换为表格的实际表头
# 或者使用索引
target_table_index = 4

# 调用方法并保存结果
# doc = process_specific_table(input_file_path, selected_items,target_table_index=target_table_index)
# doc.save(output_file_path)
#
# print(f"处理后的文档已保存至: {output_file_path}")

processor = SpecWordTableProcessor(input_file_path)
target_table_index = 4
processor.process_table(selected_items, target_table_index=target_table_index)
processor.save(input_file_path)