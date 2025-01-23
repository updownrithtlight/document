from docx import Document
from docx.shared import Inches
import os

def get_image_size(image_path):
    """获取图片宽高（英寸）"""
    from PIL import Image
    with Image.open(image_path) as img:
        width, height = img.size
        return width / 96, height / 96  # Word 默认 96 DPI

def replace_image_at_bookmark(doc, bookmark_name, image_path):
    """替换 Word 书签 `bookmark_name` 位置的图片"""
    if not os.path.exists(image_path):
        raise FileNotFoundError(f"图片 {image_path} 不存在")

    width, height = get_image_size(image_path)  # 获取原图片尺寸

    for para in doc.paragraphs:
        if bookmark_name in para.text:
            para.text = ""  # 清空占位符
            run = para.add_run()
            run.add_picture(image_path, width=Inches(width), height=Inches(height))
            break  # 替换完成后退出

# 读取 Word 文档
doc = Document("output2.docx")

# 替换 `image_placeholder` 书签为 `new_image.jpg`
replace_image_at_bookmark(doc, "image_placeholder", "new.png")

# 保存 Word 文件
doc.save("output2-2.docx")
