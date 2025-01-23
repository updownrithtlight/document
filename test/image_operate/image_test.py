from docx import Document
from docx.shared import Inches
import xml.etree.ElementTree as ET


def get_original_image_size(doc, bookmark_name):
    """
    获取 Word 文档中书签位置的图片大小
    :param doc: Word 文档对象
    :param bookmark_name: 书签名称（如 image_placeholder）
    :return: (width, height) 以英寸为单位，若找不到图片则返回默认大小
    """
    namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    for para in doc.paragraphs:
        if bookmark_name in para.text:
            # 获取段落的 XML 结构
            para_xml = ET.fromstring(para._element.xml)

            # 遍历所有图片节点
            for drawing in para_xml.findall('.//w:drawing', namespace):
                for extent in drawing.findall('.//wp:extent', {
                    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'}):
                    width = int(extent.attrib['cx']) / 12700  # EMU to inches
                    height = int(extent.attrib['cy']) / 12700  # EMU to inches
                    return width / 72, height / 72  # 转换为英寸

    return 3.5, 2.5  # 默认大小（如果没有找到图片）


def replace_image_at_bookmark(doc, bookmark_name, image_path):
    """
    在 Word 书签位置替换图片，保持原图尺寸
    :param doc: Word 文档对象
    :param bookmark_name: 书签名称（如 image_placeholder）
    :param image_path: 替换的新图片路径
    """
    # 获取原图片尺寸
    width, height = get_original_image_size(doc, bookmark_name)
    print('获取到了吗',width, height )
    # 遍历文档中的所有段落，找到书签所在段落
    for para in doc.paragraphs:
        if bookmark_name in para.text:
            para.text = ""  # 删除占位符
            run = para.add_run()
            run.add_picture(image_path, width=Inches(width), height=Inches(height))
            break  # 只替换一次，防止多次插入


# 读取 Word 文档
doc = Document("output2.docx")

# 在书签 `image_placeholder` 位置替换 `new_image.jpg`
replace_image_at_bookmark(doc, "image_placeholder", "new.png")

# 保存修改后的 Word 文件
doc.save("output.docx")
