from docx import Document


def process_section_by_marker(doc_path, marker_text, delete_section=True):
    """
    根据指定的标志段落 marker_text 定位段落，
    如果 delete_section 为 True，则删除该标志段落及其后的所有内容；
    如果 delete_section 为 False，则仅删除该标志段落，保留其后的内容。

    :param doc_path: python‑docx 的 Document 对象
    :param marker_text: 标志段落中包含的文本，用于定位该段落
    :param delete_section: 布尔值，如果为 True 删除标志段落及其之后所有内容，
                           如果为 False 仅删除标志段落
    """
    doc = Document(doc_path)

    def delete_paragraph(paragraph):
        """
        删除指定段落：直接操作底层 XML 删除 <w:p> 节点
        """
        p = paragraph._element
        p.getparent().remove(p)
        # 清空对象内部引用
        paragraph._p = paragraph._element = None

    # 遍历所有段落，查找第一个包含 marker_text 的段落
    marker_index = None
    for i, para in enumerate(doc.paragraphs):
        if marker_text in para.text:
            marker_index = i
            break

    if marker_index is None:
        print(f"未找到包含标志文本 '{marker_text}' 的段落。")
        return

    if delete_section:
        # 删除标志段落及其之后的所有段落
        # 由于删除操作会修改 doc.paragraphs，所以先复制要删除的段落列表
        paragraphs_to_delete = doc.paragraphs[marker_index:]
        for para in paragraphs_to_delete:
            delete_paragraph(para)
        print("已删除标志段落及其之后的所有内容。")
    else:
        # 仅删除标志段落，保留后续内容
        delete_paragraph(doc.paragraphs[marker_index])
        print("仅删除了标志段落，保留了其后的内容。")
    doc.save(doc_path)



# 示例：使用该方法处理 Word 文档
if __name__ == '__main__':
    # 打开 Word 文档（确保路径正确）
    doc_path = "technical_document_template.docx"

    # 定义标志文本，比如文档中有个段落包含 "### DELETE HERE ###"
    marker_text = "### DELETE HERE ###"

    # 设置是否删除标志段落后所有内容
    # 如果为 True，则删除标志段落及其之后的所有段落
    # 如果为 False，则仅删除标志段落本身
    delete_section = True  # 或 False，根据实际需求设置

    # 执行处理
    process_section_by_marker(doc_path, marker_text, delete_section)

    # 保存修改后的文档
