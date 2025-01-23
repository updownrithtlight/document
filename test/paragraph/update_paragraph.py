from docx import Document
doc = Document('technical_document_template.docx')

# 修改特定标题
for para in doc.paragraphs:
    if para.style.name.startswith("Heading") and para.text == "一级标题":
        para.text = "修改后的一级标题"

# 修改特定段落
for para in doc.paragraphs:
    if para.text == "这是第一个段落。":
        para.text = "这是修改后的第一个段落。"

doc.save('example_modified.docx')


doc = Document('example.docx')

# 在第二个段落前插入新段落
p = doc.paragraphs[1]
new_p = p.insert_paragraph_before("插入的新段落")

# 插入新的标题
doc.paragraphs[2].insert_paragraph_before("插入的新标题").style = "Heading 2"

doc.save('example_modified.docx')
