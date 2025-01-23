from docx import Document
doc = Document('technical_document_template.docx')


print("【所有标题】")
for para in doc.paragraphs:
    if para.style.name.startswith("Heading"):
        print(para.text)

print("\n【所有段落】")
for para in doc.paragraphs:
    if not para.style.name.startswith("Heading"):
        print(para.text)
