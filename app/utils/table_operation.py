from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 用于段落水平对齐


# from docx.enum.table import WD_ALIGN_VERTICAL  # 若需要垂直居中可再引入

def add_row_with_auto_serial(doc, table_index, cell_values):
    """
    在指定表格末尾添加一行：
      - 第 1 列自动生成序号 (形如 1., 2., 3.)
      - 从第 2 列开始写入 cell_values 中的数据
      - 第 1 列采用 Calibri 10pt，其余列用“宋体 小四(12pt)”。
      - 所有单元格文本均水平居中。
    """
    if table_index >= len(doc.tables):
        raise ValueError(f"表格索引 {table_index} 超出范围，当前文档仅有 {len(doc.tables)} 个表格")

    table = doc.tables[table_index]

    # 新增一行
    new_row = table.add_row()

    # 计算序号，假设表第 1 行是表头，则序号 = 当前行数 - 1
    # 如果表没有表头，想让第一行就是“1.”，可改为 len(table.rows)
    serial_number = len(table.rows) - 1
    new_row.cells[0].text = f"{serial_number}."

    # 设置第一列(序号)为 Calibri 10pt，并居中
    for paragraph in new_row.cells[0].paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 水平居中
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(10)

    # 从第二列开始填充数据，并把文本设置为“宋体 小四 + 居中”
    for col_idx in range(1, len(new_row.cells)):
        if (col_idx - 1) < len(cell_values):
            new_row.cells[col_idx].text = str(cell_values[col_idx - 1])
        else:
            new_row.cells[col_idx].text = ""

        # 给单元格里的段落和 run 设置居中、字体和大小
        _apply_font_style(new_row.cells[col_idx], font_name="宋体", font_size_pt=12)


def _apply_font_style(cell, font_name="宋体", font_size_pt=12):
    """
    给单元格内所有段落、所有 run 统一设置字体、字号和“段落居中”。
    """
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 水平居中
        for run in paragraph.runs:
            # 西文字体
            run.font.name = font_name
            # 中文字体
            run.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            # 字号
            run.font.size = Pt(font_size_pt)

    # 如果需要单元格垂直居中，也可在此设置(或在创建行后设置):
    # cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def main():
    # 1. 读取 Word 文档
    doc_path = "technical_document_template.docx"
    doc = Document(doc_path)
    # 1.	{{project_name}}	{{project_model}}	1套	粘贴标签、序列号、合格证
    # 2. 要添加的数据（不包含第一列，因为第一列自动编号）
    rows_to_add = [
        ["电源滤波组件", "MTLB32B-HNBJ-..."],
        ["保险管", "12.5A"],
        ["导电材料", "/"],
        ["出厂检测报告", "/"],
        ["试验报告", "/"],
        ["纸质合格证", "/"],
    ]

    # 3. 循环添加行
    for row_data in rows_to_add:
        add_row_with_auto_serial(doc, table_index=3, cell_values=row_data)

    # 4. 保存
    output_path = "auto_serial_result.docx"
    doc.save(output_path)
    print(f"完成，已保存至：{output_path}")


if __name__ == "__main__":
    main()
