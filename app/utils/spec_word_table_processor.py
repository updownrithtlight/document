from docx import Document


class SpecWordTableProcessor:
    def __init__(self, file_path):
        """
        初始化工具类
        :param file_path: str, 输入的 Word 文档路径
        """
        self.file_path = file_path
        self.doc = Document(file_path)

    def process_table(self, selected_items, target_table_index=None, table_header=None):
        """
        处理指定的表格
        :param selected_items: list, 已选择项目的列表
        :param target_table_index: int, 指定表格的索引（从 0 开始），默认为 None
        :param table_header: str, 用于识别目标表格的表头文本，如果提供会优先通过表头匹配表格
        :return: None
        """
        # 提取已选择的项目名称
        selected_names = [item['name'] for item in selected_items]
        # Map for replacing True/False to symbols
        boolean_map = {True: '√', False: ''}
        # 定位目标表格
        target_table = None

        if table_header:
            # 如果指定了表头，通过表头识别表格
            for table in self.doc.tables:
                if table.rows[0].cells[0].text.strip() == table_header:
                    target_table = table
                    break
        elif target_table_index is not None:
            # 如果指定了表格索引，通过索引定位表格
            target_table = self.doc.tables[target_table_index]

        if not target_table:
            raise ValueError("未找到符合条件的表格，请检查参数 `table_header` 或 `target_table_index` 是否正确。")

        # 遍历目标表格并处理数据
        rows_to_delete = []  # 保存需要删除的行
        for row_idx, row in enumerate(target_table.rows[1:], start=1):  # 跳过表头
            cell_name = row.cells[0].text.strip()
            if cell_name not in selected_names:
                # 如果不在已选择项目中，标记为删除
                rows_to_delete.append(row)
            else:
                # 如果在已选择项目中，更新内容
                for item in selected_items:
                    if item['name'] == cell_name:
                        row.cells[1].text = boolean_map[item['pcb']]
                        row.cells[2].text = boolean_map[item['beforeSeal']]
                        row.cells[3].text = boolean_map[item['afterLabel']]
                        row.cells[4].text = item['samplePlan']

                        for cell in row.cells:
                            set_cell_font(cell, font_name="Times New Roman", font_size=12)  # 小四对应 12 磅

        # 删除不需要的行
        for row in rows_to_delete:
            delete_title = row.cells[0].text
            if delete_title == "外形尺寸":
                delete_title = "外形尺寸（单位：mm）"
            delete_section_by_title(self.doc,target_title=delete_title,heading_level="Heading 2")
            target_table._element.remove(row._element)

    def save(self, output_path):
        """
        保存文档
        :param output_path: str, 输出的 Word 文档路径
        :return: None
        """
        self.doc.save(output_path)

def delete_section_by_title(doc, target_title, heading_level="Heading 1"):
    """
    删除 Word 文档中特定标题及其后续内容，直到下一个相同级别的标题。
    包括删除段落和表格。

    :param doc: Word 文档对象
    :param target_title: 要删除的标题文本
    :param heading_level: 要匹配的标题级别（默认 "Heading 1"）
    """

    def delete_element(element):
        """删除指定的元素（段落或表格）"""
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    delete_flag = False  # 标志是否处于删除范围

    # 遍历文档的 body 元素（包括段落和表格）
    for element in list(doc._element.body):
        if element.tag.endswith("p"):  # 段落
            # 获取段落对象（从 XML 元素映射回段落）
            para = next((p for p in doc.paragraphs if p._element == element), None)
            if para is None:
                continue

            if para.style.name == heading_level:  # 检查是否是标题
                if para.text.strip() == target_title:  # 找到目标标题
                    delete_flag = True
                else:
                    delete_flag = False  # 遇到下一个同级标题，停止删除

            if delete_flag:
                delete_element(element)

        elif element.tag.endswith("tbl"):  # 表格
            if delete_flag:
                delete_element(element)


from docx.shared import Pt
from docx.oxml.ns import qn
def set_cell_font(cell, font_name, font_size):
    """
    设置单元格文本的字体和字号
    :param cell: 目标单元格
    :param font_name: 字体名称，例如 "Times New Roman"
    :param font_size: 字号，例如 12（表示小四）
    """
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # 兼容中文字体
            run.font.size = Pt(font_size)


def update_table_captions_and_references(doc):
    """
    同步 Word 文档中的表格题注序号和交叉引用。
    :param doc: Document 对象
    """
    # Step 1: 查找所有表格并重新生成题注
    caption_prefix = "表"
    table_count = 0
    for paragraph in doc.paragraphs:
        if "表" in paragraph.text and "SEQ" in paragraph._element.xml:  # 查找题注字段
            table_count += 1
            # 更新题注内容
            new_caption = f"{caption_prefix} {table_count}"
            paragraph.text = new_caption

    # Step 2: 查找交叉引用并更新序号
    for paragraph in doc.paragraphs:
        if caption_prefix in paragraph.text and "REF" in paragraph._element.xml:  # 查找交叉引用字段
            # 找到并替换为正确的表序号
            ref_text = paragraph.text.split()[0]
            if ref_text.startswith(caption_prefix):
                paragraph.text = f"{caption_prefix} {table_count}"


