import os
import pythoncom
import win32com.client as win32
from docx import Document
from docxtpl import DocxTemplate

from app import logger


class WordTocTool:
    @staticmethod
    def delete_section_by_title(doc, target_title, heading_level="Heading 1"):
        """
        删除 Word 文档中，从指定标题开始直到下一个同级标题出现之前的所有段落。
        注意：此方法依赖 python‑docx 内部对象，未来版本可能需要调整。
        """

        def delete_paragraph(paragraph):
            element = paragraph._element
            element.getparent().remove(element)
            # 清空内部对象，确保段落被删除
            paragraph._p = paragraph._element = None

        delete_flag = False
        for para in doc.paragraphs[:]:
            if para.style.name == heading_level:
                if para.text.strip() == target_title:
                    delete_flag = True
                else:
                    delete_flag = False
            if delete_flag:
                delete_paragraph(para)

    @staticmethod
    def update_toc_via_word(doc_path):
        pythoncom.CoInitialize()
        word = None
        doc = None
        try:
            abs_path = os.path.abspath(doc_path)

            # 检查文件路径和权限
            if not os.path.exists(abs_path):
                logger.error(f"❌ 文件路径不存在：{abs_path}")
                return False
            if not os.access(abs_path, os.R_OK):
                logger.error(f"❌ 无法读取文件：{abs_path}")
                return False
            if not os.access(abs_path, os.W_OK):
                logger.error(f"❌ 无法写入文件：{abs_path}")
                return False

            # 检查文件是否被占用
            def is_file_locked(filepath):
                try:
                    with open(filepath, "r+"):
                        return False
                except IOError:
                    return True

            if is_file_locked(abs_path):
                logger.error(f"❌ 文件被占用：{abs_path}")
                return False

            logger.info("启动 Word 应用程序...")
            word = win32.Dispatch("Word.Application")
            word.Visible = False

            try:
                logger.info(f"尝试打开文档：{abs_path}")
                doc = word.Documents.Open(abs_path)
                if not doc:
                    logger.error(f"❌ 无法打开文档，返回值为空：{abs_path}")
                    return False
            except Exception as e:
                logger.error(f"❌ 打开文档失败：{e}")
                return False

            try:
                if doc.TablesOfContents.Count > 0:
                    logger.info("更新目录 TOC ...")
                    toc = doc.TablesOfContents(1)
                    toc.Update()
                else:
                    logger.warning("文档中没有目录 TOC")
                logger.info("保存文档...")
                doc.Save()
            except Exception as e:
                logger.error(f"❌ 更新目录失败：{e}")
                return False

            try:
                doc.Close(False)
                logger.info("关闭文档成功。")
            except Exception as e:
                logger.error(f"关闭文档时出现异常：{e}")
            finally:
                if word:
                    word.Quit()
                    logger.info("Word 应用程序已退出。")
            return True
        except Exception as e:
            logger.error(f"❌ 未知错误：{e}", exc_info=True)
            return False
        finally:
            pythoncom.CoUninitialize()

    @staticmethod
    def delete_section_by_title2_or_higher(doc, target_title):
        """
        删除 Word 文档中，从指定标题2（Heading 2）开始，直到遇到下一个相同级别 (Heading 2) 或更高级别 (Heading 1) 的所有段落。

        :param doc: Word 文档对象 (Document)
        :param target_title: 需要删除的标题2（Heading 2）的文本
        """

        def delete_paragraph(paragraph):
            """
            删除段落的底层 XML 结构
            """
            element = paragraph._element
            element.getparent().remove(element)
            paragraph._p = paragraph._element = None

        delete_flag = False  # 是否进入删除模式

        for para in doc.paragraphs[:]:  # 遍历所有段落
            para_text = para.text.strip()  # 获取文本内容
            para_style = para.style.name  # 获取样式名称

            # 遇到 `Heading 1` 或 `Heading 2`，可能要停止删除
            if para_style in ["Heading 1", "Heading 2"]:
                if delete_flag:
                    # 如果已经进入删除模式，遇到 Heading 1 或 Heading 2 就终止删除
                    break

                # 如果当前段落是目标标题 (Heading 2)，则开始删除模式
                if para_text == target_title and para_style == "Heading 2":
                    delete_flag = True  # 开启删除模式
                    delete_paragraph(para)  # 删除目标标题本身

            elif delete_flag:
                # 处于删除模式，则删除该段落
                delete_paragraph(para)

    @staticmethod
    def process_document(template_path, new_doc_suffix="_modified", delete_titles=None):
        """
        加载模板文档后，删除指定标题（列表 delete_titles 中的每个标题）及其后续内容，
        然后保存为新文档，并返回新文档路径。

        :param template_path: 模板文档的路径（.docx 文件）。
        :param new_doc_suffix: 新文档文件名后缀（默认 "_modified"）。
        :param delete_titles: 需要删除章节的标题列表，例如 ["插入损耗特性", "产品功能"]。
        :return: 新文档的路径，若处理失败则返回 None。
        """
        if not os.path.exists(template_path):
            print("❌ 模板文件不存在：", template_path)
            return None

        new_doc_path = template_path.replace(".docx", f"{new_doc_suffix}.docx")
        print("加载模板文档...")
        try:
            doc = Document(template_path)
        except Exception as e:
            print("❌ 加载文档失败：", e)
            return None

        if delete_titles:
            for title in delete_titles:
                print(f"删除标题 '{title}' 及其后续内容...")
                WordTocTool.delete_section_by_title(doc, title)

        try:
            doc.save(new_doc_path)
            print(f"✅ 文档修改成功，已保存到：{new_doc_path}")
        except Exception as e:
            print("❌ 保存文档失败：", e)
            return None

        return new_doc_path

    @staticmethod
    def fill_doc_with_features(template_path, context):
        # 载入模板
        doc = DocxTemplate(template_path)

        # 渲染并保存
        doc.render(context)
        doc.save(template_path)



if __name__ == "__main__":
    # 示例：请根据实际情况修改模板文件路径和需要删除的标题列表
    template_path = r"C:\noneSystem\bj\document\test\technical_document_template.docx"
    delete_titles = ['环境特性']

    # 先处理文档：删除指定章节，并保存为新文件
    new_doc_path = WordTocTool.process_document(template_path, new_doc_suffix="_modified", delete_titles=delete_titles)

    # 如果文档处理成功，则调用 COM 自动化更新目录
    if new_doc_path:
        if WordTocTool.update_toc_via_word(new_doc_path):
            print("✅ TOC 更新成功")
        else:
            print("❌ TOC 更新失败")
